from __future__ import annotations

from typing import BinaryIO
from typing import Iterator

import docx

from .base_parser import BaseParser


class DocxParser(BaseParser):
    """Document parser for DOCX files using python-docx."""

    def parse_file(self, file_stream: BinaryIO) -> str:
        """
        Parse the entire DOCX document stream into a single string.
        """
        # python-docx can accept a file-like object directly
        doc = docx.Document(file_stream)
        text = [paragraph.text for paragraph in doc.paragraphs if paragraph.text]
        return "\n".join(text)

    def parse_file_stream(self, file_stream: BinaryIO) -> Iterator[str]:
        """
        Parse the DOCX document stream paragraph-by-paragraph.
        """
        doc = docx.Document(file_stream)
        for paragraph in doc.paragraphs:
            if paragraph.text:
                yield paragraph.text
